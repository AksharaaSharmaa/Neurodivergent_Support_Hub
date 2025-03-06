import streamlit as st
import faiss
import numpy as np
import json
import PyPDF2
import docx
import io
import google.generativeai as genai
import tensorflow_hub as hub
import tensorflow as tf
from langchain.text_splitter import RecursiveCharacterTextSplitter
from gtts import gTTS
import tempfile
import os
from datetime import datetime
from typing import List, Dict, Tuple
import re
from sklearn.metrics.pairwise import cosine_similarity
import base64
import pyttsx3
import soundfile as sf
import speech_recognition as sr
from audio_recorder_streamlit import audio_recorder
import wave
import warnings
from google.cloud import secretmanager

# Back button
if st.button("← Back to Home"):
    st.switch_page("HomePage.py")

# Ignore all warnings
warnings.filterwarnings("ignore")

# Function to access secrets from Google Secret Manager
def access_secret(secret_name):
    client = secretmanager.SecretManagerServiceClient()
    project_id = "neurodivai"
    secret_path = f"projects/{project_id}/secrets/{secret_name}/versions/latest"
    response = client.access_secret_version(request={"name": secret_path})
    return response.payload.data.decode("UTF-8")

# Configure Gemini API securely
GEMINI_API_KEY = access_secret("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

# --------------- Constants ---------------
NEURODIVERGENT_CONDITIONS = {
    "ADHD": [
        "Use time-blocking techniques for task management",
        "Break large tasks into smaller, manageable chunks",
        "Implement the Pomodoro Technique for focused work",
        "Utilize visual task management tools and boards",
        "Minimize workplace distractions with noise-canceling headphones",
        "Take regular movement breaks (every 30-60 minutes)",
        "Maintain a structured workspace with clear organization systems",
        "Use digital reminders and calendar alerts"
    ],
    
    "Autism": [
        "Establish clear, consistent daily routines and expectations",
        "Create quiet spaces for focused work and decompression",
        "Use written communication for complex instructions",
        "Implement sensory-friendly lighting and noise reduction",
        "Create detailed process documents and checklists",
        "Allow flexible work hours to accommodate sensory needs",
        "Use explicit, clear communication without idioms",
        "Provide advance notice for meetings and changes"
    ],
    
    "Dyslexia": [
        "Utilize text-to-speech software for reading tasks",
        "Implement color-coding systems for file organization",
        "Provide information in multiple formats (text, audio, visual)",
        "Allow extra time for reading and writing tasks",
        "Use dyslexia-friendly fonts in documents",
        "Encourage mind mapping for project planning",
        "Use speech-to-text tools for documentation",
        "Provide access to spell-checking and grammar tools"
    ],
    
    "Dyscalculia": [
        "Use calculators and digital tools for calculations",
        "Implement double-checking systems for numerical work",
        "Provide step-by-step written procedures",
        "Create visual aids and diagrams for numerical concepts",
        "Allow extra time for financial or numerical tasks",
        "Utilize spreadsheet templates with built-in formulas",
        "Use digital tools for time and budget tracking",
        "Implement color-coding for numerical categories"
    ],
    
    "Dyspraxia": [
        "Organize workspace to minimize physical obstacles",
        "Use ergonomic equipment and accessories",
        "Provide clear physical workspace boundaries",
        "Implement digital organization tools",
        "Allow extra time for physical tasks",
        "Create detailed checklists for sequential activities",
        "Use voice recognition software when needed",
        "Ensure adequate space between furniture and equipment"
    ],
    
    "Tourette's Syndrome": [
        "Offer flexible break schedules for self-regulation",
        "Provide private workspace options when needed",
        "Implement stress-reduction techniques",
        "Allow remote work options during high-stress periods",
        "Create an accepting and informed workplace culture",
        "Ensure access to quiet spaces",
        "Develop clear communication protocols",
        "Allow for movement and stretching breaks"
    ],
    
    "OCD": [
        "Create clear, consistent workplace protocols",
        "Allow for personalized organization systems",
        "Provide structured daily schedules",
        "Implement verification systems for tasks",
        "Allow time for preferred organizational methods",
        "Maintain consistent workplace routines",
        "Provide digital tools for task tracking",
        "Create backup systems for important work"
    ],
    
    "Bipolar Disorder": [
        "Implement flexible scheduling options",
        "Create predictable daily routines",
        "Provide quiet spaces for stress management",
        "Allow work-from-home options when needed",
        "Maintain open communication channels",
        "Develop crisis management plans",
        "Structure work around energy levels",
        "Enable adjustable workload management"
    ],
    
    "Sensory Processing Disorder": [
        "Customize lighting and noise levels in workspace",
        "Allow use of noise-canceling headphones",
        "Provide alternative seating options",
        "Create designated quiet workspaces",
        "Permit use of sensory aids and tools",
        "Implement flexible break schedules",
        "Allow for workspace modifications",
        "Minimize unexpected sensory disruptions"
    ],
    
    "Other": [
        "Develop personalized accommodation plans",
        "Maintain open communication about needs",
        "Implement flexible work arrangements",
        "Provide necessary tools and resources",
        "Create an inclusive work environment",
        "Allow for regular feedback and adjustments",
        "Establish clear support systems",
        "Enable access to workplace mentoring"
    ]
}

# --------------- Document Classes ---------------
class Document:
    def __init__(self, content: str, metadata: Dict):
        self.content = content
        self.metadata = metadata
        self.chunks = []
        self.embeddings = []
    
    def __repr__(self):
        return f"Document(source={self.metadata.get('source')}, chunks={len(self.chunks)})"

class DocumentChunk:
    def __init__(self, text: str, metadata: Dict, embedding: np.ndarray = None):
        self.text = text
        self.metadata = metadata
        self.embedding = embedding
    
    def __repr__(self):
        return f"Chunk(length={len(self.text)}, source={self.metadata.get('source')})"

# --------------- Document Processing Functions ---------------
def clean_text(text: str) -> str:
    """Clean and normalize text content."""
    text = re.sub(r'\s+', ' ', text.strip())
    text = re.sub(r'[^\w\s.,!?-]', '', text)
    return text

def extract_text_with_layout(file) -> Tuple[str, Dict]:
    """Enhanced text extraction with layout information."""
    metadata = {
        'source': file.name,
        'timestamp': datetime.now().isoformat(),
        'file_type': file.type,
        'size': file.size
    }
    
    if file.type == "text/plain":
        content = file.read().decode("utf-8")
        metadata['structure'] = 'plain_text'
        
    elif file.type == "application/pdf":
        reader = PyPDF2.PdfReader(file)
        content = ""
        metadata['pages'] = len(reader.pages)
        metadata['structure'] = 'pdf'
        
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                content += f"\nPage {i+1}:\n{page_text}"
                
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        content = ""
        metadata['structure'] = 'docx'
        metadata['sections'] = len(doc.sections)
        
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name
                content += f"\n{style}: {para.text}"
                
    return clean_text(content), metadata

def generate_audio_bytes(text: str) -> bytes:
    """
    Generates audio bytes from text using pyttsx3.
    """
    try:
        # Initialize the TTS engine
        engine = pyttsx3.init()
        
        # Create an in-memory bytes buffer
        audio_bytes_io = io.BytesIO()
        
        # Save the audio to the buffer
        engine.save_to_file(text, 'temp.wav')
        engine.runAndWait()
        
        # Read the saved file and convert to bytes
        with open('temp.wav', 'rb') as audio_file:
            audio_bytes = audio_file.read()
            
        # Clean up the temporary file
        os.remove('temp.wav')
        
        return audio_bytes
    except Exception as e:
        st.error(f"Error generating audio: {str(e)}")
        return None
    
def process_audio(audio_bytes):
    """Process audio with detailed error handling and debug info"""
    try:
        # Create a temporary WAV file in memory with specific parameters
        wav_buffer = io.BytesIO()
        with wave.open(wav_buffer, 'wb') as wav_file:
            wav_file.setnchannels(1)  # Mono
            wav_file.setsampwidth(2)  # 2 bytes per sample
            wav_file.setframerate(16000)  # Common sample rate for speech recognition
            wav_file.writeframes(audio_bytes)
        
        wav_buffer.seek(0)  # Rewind the buffer
        
        # Initialize recognizer with adjusted parameters
        recognizer = sr.Recognizer()
        recognizer.energy_threshold = 300  # Adjust for sensitivity
        recognizer.dynamic_energy_threshold = True
        recognizer.pause_threshold = 0.8  # Adjust for shorter pauses
        
        # Process the audio
        with sr.AudioFile(wav_buffer) as source:
            st.info("Processing audio...")
            # Adjust for ambient noise
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            # Record the entire file
            audio_data = recognizer.record(source)
            
            # Try different speech recognition services
            try:
                st.info("Attempting speech recognition...")
                text = recognizer.recognize_google(audio_data, language='en-US')
                st.success("Audio processed successfully!")
                return text
            except sr.UnknownValueError:
                st.error("Could not understand the audio. Please try:")
                st.write("- Speaking more clearly and closer to the microphone")
                st.write("- Reducing background noise")
                st.write("- Recording in a quieter environment")
                return None
            except sr.RequestError as e:
                st.error(f"Could not reach Google Speech Recognition service: {str(e)}")
                return None
                
    except Exception as e:
        st.error(f"Error processing audio: {str(e)}")
        st.write("Debug info:")
        st.write(f"- Audio bytes length: {len(audio_bytes)}")
        st.write(f"- Error type: {type(e).__name__}")
        return None

# Main recording interface
def audio_input_section():
    """Handle audio recording and processing"""
    audio_bytes = audio_recorder(
        text="🎤 Click to record",
        recording_color="#e8b62c",
        neutral_color="#6aa36f",
        key="audio_recorder",
        pause_threshold=2.0,  # Longer pause threshold
        sample_rate=16000  # Adjusted sample rate
    )
    
    if audio_bytes:
        st.info("Audio recorded! Processing...")
        text = process_audio(audio_bytes)
        if text:
            st.session_state.temp_text = text
            
# --------------- Vector Store Class ---------------
class EnhancedVectorStore:
    def __init__(self, embedding_dim: int):
        self.index = faiss.IndexFlatL2(embedding_dim)
        self.documents: List[Document] = []
        self.chunks: List[DocumentChunk] = []
        
    def add_document(self, document: Document):
        self.documents.append(document)
        self.chunks.extend(document.chunks)
        embeddings = np.array([chunk.embedding for chunk in document.chunks])
        self.index.add(embeddings)
        
    def search(self, query_embedding: np.ndarray, k: int = 3) -> List[Tuple[DocumentChunk, float]]:
        distances, indices = self.index.search(query_embedding.reshape(1, -1), k)
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(self.chunks):
                chunk = self.chunks[idx]
                score = 1 / (1 + distances[0][i])
                results.append((chunk, score))
        return results

# --------------- RAG Processor Class ---------------
class RAGProcessor:
    def __init__(self, embed_model):
        self.embed_model = embed_model
        self.vector_store = EnhancedVectorStore(512)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", ".", "!", "?", ","]
        )
    
    def process_document(self, file) -> Document:
        content, metadata = extract_text_with_layout(file)
        document = Document(content, metadata)
        texts = self.text_splitter.split_text(content)
        
        for i, text in enumerate(texts):
            chunk_metadata = {
                **metadata,
                'chunk_index': i,
                'chunk_size': len(text)
            }
            embedding = self.embed_model([text])[0].numpy()
            chunk = DocumentChunk(text, chunk_metadata, embedding)
            document.chunks.append(chunk)
            document.embeddings.append(embedding)
        
        self.vector_store.add_document(document)
        return document

    def generate_response(self, query: str, condition: str, support_type: str, strengths: str) -> str:
        """Generate response using Gemini, incorporating user's workplace strengths."""
        try:
            model = genai.GenerativeModel("gemini-1.5-flash")
            
            # Generate base prompt with strengths included
            prompt = f"""As a workplace support specialist with expertise in neurodivergent conditions, 
            particularly {condition}, provide detailed advice for the following situation.
            
            Support type requested: {support_type}
            
            User's Workplace Strengths: {strengths}
            
            User Query: {query}
            
            Please provide specific, practical advice that takes into account both:
            1. The unique characteristics and needs of someone with {condition}
            2. The user's stated workplace strengths and how to leverage them
            
            Include actionable steps and accommodation suggestions where appropriate, 
            and specifically mention how the user's strengths can be utilized to address their challenges."""
            
            # [Rest of the method remains the same]
            
            if len(self.vector_store.chunks) > 0:
                query_embedding = self.embed_model([query])[0].numpy()
                relevant_chunks = self.vector_store.search(query_embedding)
                relevant_chunks = [(chunk, score) for chunk, score in relevant_chunks if score > 0.5]
                
                if relevant_chunks:
                    context = "\n\n".join([
                        f"[Relevance: {score:.2f}] {chunk.text}" 
                        for chunk, score in relevant_chunks
                    ])
                    prompt += f"\n\nAdditional context from documents:\n{context}"
            
            response = model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            st.error(f"Error generating response: {str(e)}")
            return "I apologize, but I encountered an error generating the response. Please try again."


# --------------- Initialize Models ---------------
@st.cache_resource
def initialize_models():
    embed_model = hub.load("https://tfhub.dev/google/universal-sentence-encoder/4")
    return RAGProcessor(embed_model)

def audio_input_section(session_state_key: str):
    """Handle audio recording and processing with specified session state key"""
    audio_bytes = audio_recorder(
        text="🎤 Click to record",
        recording_color="#e8b62c",
        neutral_color="#6aa36f",
        key=f"audio_recorder_{session_state_key}",  # Unique key for each recorder
        pause_threshold=2.0,
        sample_rate=16000
    )
    
    if audio_bytes:
        st.info("Audio recorded! Processing...")
        text = process_audio(audio_bytes)
        if text:
            setattr(st.session_state, session_state_key, text)

# --------------- Streamlit UI ---------------
def main():
    # Initialize session state if it doesn't exist
    if 'temp_text' not in st.session_state:
        st.session_state.temp_text = ""
    if 'temp_strengths' not in st.session_state:
        st.session_state.temp_strengths = ""

    # Initialize RAG processor
    rag_processor = initialize_models()

    # Custom CSS for consistent formatting
    st.markdown("""
        <style>
        .main-header {
            font-size: 2.5rem;
            font-weight: bold;
            margin-bottom: 2rem;
        }
        .section-header {
            font-size: 1.8rem;
            font-weight: bold;
            margin: 1.5rem 0;
        }
        .subsection-header {
            font-size: 1.3rem;
            font-weight: bold;
            margin: 1rem 0;
        }
        .stSelectbox, .stTextArea {
            margin: 1rem 0;
        }
        .sidebar-header {
            font-size: 1.5rem;
            font-weight: bold;
            margin: 1rem 0;
        }
        </style>
    """, unsafe_allow_html=True)

    # Main Title
    st.title('🧠 Neurodivergent Workplace Support')

    # Sidebar Configuration
    with st.sidebar:
        st.markdown('<p class="sidebar-header">📂 Reference Documents</p>', unsafe_allow_html=True)
        st.markdown("Upload supporting documents for enhanced responses")
        uploaded_files = st.file_uploader(
            "Upload files (TXT, PDF, DOCX)",
            type=["txt", "pdf", "docx"],
            accept_multiple_files=True,
            help="Your documents will be processed to provide more relevant responses"
        )

        if uploaded_files:
            for file in uploaded_files:
                with st.spinner(f"Processing {file.name}..."):
                    document = rag_processor.process_document(file)
                    st.success(f"Processed {file.name}: {len(document.chunks)} chunks created")

    # Main Form Container
    with st.container():
        st.markdown('<p class="section-header">🤖 Personal Support Configuration</p>', unsafe_allow_html=True)
        
        # Two-column layout for main selections
        col1, col2 = st.columns(2)
        with col1:
            selected_condition = st.selectbox(
                "Neurodivergent Condition",
                list(NEURODIVERGENT_CONDITIONS.keys()),
                format_func=lambda x: f"{x} - {NEURODIVERGENT_CONDITIONS[x]}"
            )
        
        with col2:
            support_options = [
                "Personalized Workplace Advice",
                "Communication Script",
                "Email Template for Accommodation Requests",
                "Workplace Strategy Recommendations",
                "Conflict Resolution Guidance",
                "Task Management Strategies",
                "Sensory Environment Adjustments",
                "Meeting Participation Strategies"
            ]
            selected_support = st.selectbox("Type of Support Needed", support_options)

    # Strengths Input Section
    st.markdown('<p class="subsection-header">💪 Your Workplace Strengths</p>', unsafe_allow_html=True)
    strengths_col1, strengths_col2 = st.columns([3, 1])
    
    with strengths_col1:
        strengths = st.text_area(
            "List your workplace strengths",
            value=st.session_state.temp_strengths,
            placeholder="Example: Detail-oriented, creative problem-solver, strong analytical skills...",
            height=100
        )
    
    with strengths_col2:
        st.markdown("##### Voice Input")
        audio_input_section("temp_strengths")

    # Query Input Section
    st.markdown('<p class="subsection-header">❓ Your Situation</p>', unsafe_allow_html=True)
    query_col1, query_col2 = st.columns([3, 1])
    
    with query_col1:
        query = st.text_area(
            "Describe your workplace situation or concern",
            value=st.session_state.temp_text,
            height=150,
            placeholder="Be specific about your workplace challenge or question..."
        )
    
    with query_col2:
        st.markdown("##### Voice Input")
        audio_input_section("temp_text")

    # Response Options
    st.markdown('<p class="subsection-header">🎯 Response Preferences</p>', unsafe_allow_html=True)
    enable_voice = st.checkbox(
        "Enable Voice Support",
        help="Generate an audio version of the response"
    )

    # Generate Response Button
    if st.button("Generate Response", type="primary"):
        if query:
            with st.spinner("Generating personalized response..."):
                response = rag_processor.generate_response(
                    query=query,
                    condition=selected_condition,
                    support_type=selected_support,
                    strengths=strengths
                )
                
                # Response Display
                st.markdown('<p class="section-header">📝 Your Personalized Response</p>', unsafe_allow_html=True)
                st.markdown(response)
                
                # Download Options
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "📥 Download as Text",
                        response,
                        file_name=f"{selected_condition}_{selected_support}_response.txt",
                        mime="text/plain"
                    )
                
                if enable_voice:
                    with col2:
                        st.markdown("### 🔊 Audio Version")
                        audio_bytes = generate_audio_bytes(response)
                        if audio_bytes:
                            st.audio(audio_bytes, format='audio/wav')
        else:
            st.warning("Please enter a query or use voice input.")

    # Condition-specific tips in sidebar
    if selected_condition != "Other":
        with st.sidebar:
            st.markdown(f'<p class="sidebar-header">💡 Workplace Management Strategies for {selected_condition}</p>', 
                    unsafe_allow_html=True)
            strategies = NEURODIVERGENT_CONDITIONS[selected_condition]
            for strategy in strategies:
                st.markdown(f"* {strategy}")
                
if __name__ == "__main__":
    main()